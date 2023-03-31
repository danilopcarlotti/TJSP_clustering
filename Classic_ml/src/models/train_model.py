from collections import Counter
from docx import Document
from pathlib import Path
from sklearn.cluster import KMeans
from sklearn.metrics import pairwise_distances_argmin_min
from sklearn.metrics.cluster import homogeneity_score
from sklearn.neighbors import KDTree
from sklearn.preprocessing import MinMaxScaler

import numpy as np
import re
import sys

sys.path.append(str(Path().absolute().parent))
try:
    from data.make_dataset import load_df
    from features.build_features import hashing_texts
except:
    sys.path.append(str(Path().absolute().parent.parent))
    from src.data.make_dataset import load_df
    from src.features.build_features import hashing_texts


def cluster_documents(args: dict, n_clusters: int = 15):
    scaler = MinMaxScaler()
    report = {}
    print("Preprocessing data")
    df = load_df(args["path_file"], args["type_file"])
    if int(args["just_final_decisions"]):
        texts = [
            i
            for i in df["text"].tolist()
            if re.search(
                r"julgo.{,15}procedente|decido\.|acordam|conhec.{,10}recurso", i, re.I
            )
        ][:15000]
    else:
        texts = df["text"].tolist()[:15000]
    report["Número total de textos analisados"] = len(df.index)
    y = df["class"].tolist()
    X = hashing_texts(texts)
    X = scaler.fit_transform(X)
    print("Clustering with kmeans")
    if len(X) < 15:
        report["Kmeans_inviavel"] = 1
    else:
        report["Kmeans_inviavel"] = 0
        kmeans_model = KMeans(n_clusters=n_clusters, random_state=0).fit(X)
        labels = kmeans_model.labels_
        counter_kmeans = dict(Counter(labels))
        report["Número kmeans"] = len(counter_kmeans.keys())
        for kmeans_label, count_label in counter_kmeans.items():
            vectors_class = X[
                np.ix_(
                    [
                        index
                        for index, label in enumerate(labels)
                        if label == kmeans_label
                    ]
                )
            ]
            mean_vectors = np.mean(vectors_class, axis=0)
            tree = KDTree(vectors_class, leaf_size=10)
            _, indexes_centroids = tree.query([mean_vectors])
            report["Texto representativo da classe {}".format(kmeans_label)] = texts[
                indexes_centroids[0][0]
            ]
            report["estatísticas_textos_{}".format(kmeans_label)] = "{:.2f}%".format(
                100 * (count_label / len(df.index))
            )

    if len(y):
        hom_score = homogeneity_score(y, kmeans_model.labels_)
    else:
        hom_score = -1
    counter_y = dict(Counter(y))
    report["classes"] = counter_y.keys()
    for k, v in counter_y.items():
        report["number_texts_class_{}".format(k)] = "{:.2f}%".format(
            100 * (v / len(df.index))
        )
    report["homogeneity_score"] = hom_score
    return report


def save_report_docx(
    report: dict, output_path: str = "", name_report: str = "relatório"
):
    text_file = open(output_path + "/{}.txt".format(name_report), "w", encoding="utf-8")
    text_file.write(
        "Número total de textos analisados: {}\n\n".format(
            report["Número total de textos analisados"]
        )
    )
    for c in report["classes"]:
        text_file.write(
            "Percentual de textos da classe {} : {}\n".format(
                c, report["number_texts_class_{}".format(c)]
            )
        )
    text_file.write(
        "\n\nScore de homogeneidade: {}\n\n".format(report["homogeneity_score"])
    )
    for k in range(report["Número kmeans"]):
        text_file.write(
            "\n\nTexto representativo da classe {} que representa {} dos textos:\n {}\n\n".format(
                str(k),
                str(report["estatísticas_textos_{}".format(k)]),
                re.sub(
                    r"\s+",
                    " ",
                    str(report["Texto representativo da classe {}".format(k)]),
                ),
            )
        )
        text_file.write(80 * "*")
        text_file.write("\n\n")
    text_file.close()
    # doc = Document()
    # doc.add_paragraph(
    #     "Número total de textos analisados: {}".format(
    #         report["Número total de textos analisados"]
    #     )
    # )
    # doc.add_paragraph(
    #     "\nScore de homogeneidade: {}\n".format(report["homogeneity_score"])
    # )
    # for k in range(report["Número kmeans"]):
    #     doc.add_paragraph(
    #         "Texto representativo da classe {} que representa {} dos textos: {}".format(
    #             str(k),
    #             str(report["estatísticas_textos_{}".format(k)]),
    #             str(report["Texto representativo da classe {}".format(k)]),
    #         )
    #     )
    #     doc.add_paragraph(80 * "*")
    #     doc.add_paragraph("\n")
    # doc.save(output_path + "{}.docx".format(name_report))


if __name__ == "__main__":
    pass
